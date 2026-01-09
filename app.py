# app.py
# --------------------------------------------------
# Turkish Market Newsletter – Streamlit Dashboard
# Fully audited & cleaned version (v1.0)
# --------------------------------------------------

import streamlit as st
import feedparser
import yfinance as yf
from datetime import datetime, timedelta
from docx import Document
from fpdf import FPDF
import os

# ==================================================
# PAGE CONFIG
# ==================================================

st.set_page_config(
    page_title="Turkish Market Newsletter",
    layout="wide"
)

# ==================================================
# CONSTANTS & CONFIG
# ==================================================

RSS_FEEDS = {
    "Bigpara": "https://www.bigpara.com/rss/",
    "BloombergHT": "https://www.bloomberght.com/rss",
    "ReutersTR": "https://feeds.reuters.com/reuters/TurkeyNews"
}

KEYWORDS = {
    "high": ["TCMB", "faiz", "enflasyon", "Fed", "ECB", "bilanço", "KAP"],
    "medium": ["BIST", "endeks", "döviz", "CDS", "rezerv", "petrol"],
}

OUTPUT_DIR = "output"
os.makedirs(OUTPUT_DIR, exist_ok=True)

TODAY_STR = datetime.now().strftime("%d %B %Y")

# ==================================================
# DATA COLLECTION (CACHED)
# ==================================================

@st.cache_data(ttl=900)
def fetch_news():
    articles = []
    cutoff = datetime.now() - timedelta(hours=18)

    for source, url in RSS_FEEDS.items():
        feed = feedparser.parse(url)
        for entry in feed.entries:
            published = (
                datetime(*entry.published_parsed[:6])
                if hasattr(entry, "published_parsed")
                else datetime.now()
            )
            if published >= cutoff:
                articles.append({
                    "title": entry.title,
                    "summary": entry.get("summary", ""),
                    "source": source,
                    "published": published
                })
    return articles


@st.cache_data(ttl=900)
def get_market_data():
    bist = yf.Ticker("XU100.IS").history(period="1d")
    usdtry = yf.Ticker("USDTRY=X").history(period="1d")

    return {
        "bist_close": round(float(bist["Close"].iloc[-1]), 2),
        "bist_change": round((bist["Close"].iloc[-1] / bist["Open"].iloc[-1] - 1) * 100, 2),
        "usdtry": round(float(usdtry["Close"].iloc[-1]), 2)
    }

# ==================================================
# NEWS SCORING & SELECTION
# ==================================================

def score_news(article):
    score = 0
    text = (article["title"] + " " + article["summary"]).lower()

    for kw in KEYWORDS["high"]:
        if kw.lower() in text:
            score += 3

    for kw in KEYWORDS["medium"]:
        if kw.lower() in text:
            score += 1

    if article["source"] in ["ReutersTR", "BloombergHT"]:
        score += 2

    hours_old = (datetime.now() - article["published"]).seconds / 3600
    score += max(0, 3 - hours_old)

    return score


def select_top_news(articles, n=3):
    for a in articles:
        a["score"] = score_news(a)
    articles.sort(key=lambda x: x["score"], reverse=True)
    return articles[:n]

# ==================================================
# MARKET INTELLIGENCE LAYERS
# ==================================================

def generate_market_summary(market):
    direction = "yükselişle" if market["bist_change"] > 0 else "düşüşle"
    return (
        f"BIST 100 günü %{abs(market['bist_change'])} {direction} "
        f"{market['bist_close']} seviyesinde tamamladı. "
        f"USD/TRY {market['usdtry']} seviyesinde izleniyor."
    )


def why_this_matters(news):
    title = news["title"].lower()

    if any(k in title for k in ["faiz", "tcmb", "merkez bankası"]):
        return "Para politikası adımları, özellikle bankacılık sektörü olmak üzere tüm piyasa değerlemelerini etkiler."

    if any(k in title for k in ["bilanço", "kar", "zarar"]):
        return "Finansal sonuçlar, şirketin operasyonel gücünü ve mevcut fiyatlamaların sürdürülebilirliğini gösterir."

    if any(k in title for k in ["fed", "abd", "enflasyon"]):
        return "Küresel makro gelişmeler, gelişen piyasalara yönelik risk iştahını ve sermaye akımlarını belirler."

    if any(k in title for k in ["petrol", "emtia", "altın"]):
        return "Emtia fiyatları, enflasyon beklentileri ve ilgili sektörler üzerinde belirleyici rol oynar."

    return "Bu gelişme, yatırımcı algısı ve piyasa beklentileri açısından önem taşıyor."


def sector_impact(news):
    text = (news["title"] + " " + news["summary"]).lower()
    sectors = []

    if any(k in text for k in ["faiz", "tcmb", "banka", "kredi", "mevduat"]):
        sectors.append("Banking")
    if any(k in text for k in ["sanayi", "üretim", "ihracat", "fabrika"]):
        sectors.append("Industrial")
    if any(k in text for k in ["enerji", "petrol", "doğalgaz", "elektrik"]):
        sectors.append("Energy")

    return sectors or ["Broad Market"]


def sector_heat(top_news):
    heat = {"Banking": 0, "Industrial": 0, "Energy": 0}
    pos = ["artış", "yükseliş", "güçlü", "rekor", "olumlu"]
    neg = ["düşüş", "gerileme", "zayıf", "baskı", "risk"]

    for n in top_news:
        text = (n["title"] + " " + n["summary"]).lower()
        sentiment = (1 if any(p in text for p in pos) else 0) - (1 if any(m in text for m in neg) else 0)
        for s in sector_impact(n):
            if s in heat:
                heat[s] += sentiment

    labels = {}
    for k, v in heat.items():
        labels[k] = "🔥 Positive" if v > 0 else "❄️ Negative" if v < 0 else "➖ Neutral"
    return labels

# ==================================================
# EXPORT FUNCTIONS
# ==================================================

def export_docx(top_news, summary):
    doc = Document()
    doc.add_heading(f"Günlük Piyasa Bülteni – {TODAY_STR}", level=1)

    for i, n in enumerate(top_news, 1):
        doc.add_heading(f"{i}. {n['title']}", level=2)
        doc.add_paragraph(n["summary"])
        doc.add_paragraph(f"Why this matters: {why_this_matters(n)}")

    doc.add_heading("Piyasa Özeti", level=2)
    doc.add_paragraph(summary)

    path = f"{OUTPUT_DIR}/newsletter.docx"
    doc.save(path)
    return path


def export_pdf(top_news, summary):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    pdf.multi_cell(0, 8, f"Günlük Piyasa Bülteni – {TODAY_STR}\n\n")

    for i, n in enumerate(top_news, 1):
        pdf.multi_cell(0, 8, f"{i}. {n['title']}\n{n['summary']}\nWhy this matters: {why_this_matters(n)}\n")

    pdf.multi_cell(0, 8, f"\nPiyasa Özeti\n{summary}")

    path = f"{OUTPUT_DIR}/newsletter.pdf"
    pdf.output(path)
    return path

# ==================================================
# STREAMLIT UI
# ==================================================

st.title("📈 Turkish Market Daily Newsletter")
st.caption("Noise-free, sector-aware market intelligence")

if st.button("🔄 Fetch Today's Data"):
    with st.spinner("Loading data..."):
        news = fetch_news()
        top_news = select_top_news(news)
        market = get_market_data()
        summary = generate_market_summary(market)
        heat = sector_heat(top_news)

    st.subheader("📊 Market Snapshot")
    c1, c2, c3 = st.columns(3)
    c1.metric("BIST 100", market["bist_close"], f"{market['bist_change']}%")
    c2.metric("USD/TRY", market["usdtry"])
    c3.metric("Market Mood", "Positive" if market["bist_change"] > 0 else "Negative")

    st.subheader("🌡️ Sector Heat Indicator")
    h1, h2, h3 = st.columns(3)
    h1.metric("🏦 Banking", heat["Banking"])
    h2.metric("🏭 Industrial", heat["Industrial"])
    h3.metric("⚡ Energy", heat["Energy"])

    st.subheader("📰 Top 3 News")
    for i, n in enumerate(top_news, 1):
        with st.expander(f"{i}. {n['title']}"):
            st.write(n["summary"])
            st.markdown(f"**Sector Impact:** {', '.join(sector_impact(n))}")
            st.markdown("**Why this matters:**")
            st.write(why_this_matters(n))
            st.caption(f"Source: {n['source']} | Score: {round(n['score'],2)}")

    st.subheader("📄 Export")
    if st.button("Generate Word & PDF"):
        docx_path = export_docx(top_news, summary)
        pdf_path = export_pdf(top_news, summary)

        st.success("Files generated")
        st.download_button("⬇️ Download DOCX", open(docx_path, "rb"), file_name="newsletter.docx")
        st.download_button("⬇️ Download PDF", open(pdf_path, "rb"), file_name="newsletter.pdf")
